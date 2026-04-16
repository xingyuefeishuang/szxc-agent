import sys
import os
import json
from docx import Document

def read_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found.")
        sys.exit(1)
    
    doc = Document(file_path)
    result = {
        "paragraphs": [],
        "tables": [],
        "core_properties": {}
    }

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            result["paragraphs"].append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result["tables"].append(table_data)

    # Extract metadata
    cp = doc.core_properties
    result["core_properties"] = {
        "author": cp.author,
        "created": str(cp.created),
        "last_modified_by": cp.last_modified_by,
        "modified": str(cp.modified),
        "title": cp.title
    }

    return result

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    try:
        data = read_docx(file_path)
        print(json.dumps(data, ensure_ascii=False, indent=2))
    except Exception as e:
        print(f"Error reading docx: {str(e)}")
        sys.exit(1)
